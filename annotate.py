from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_paragraph_bg_color(paragraph, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    paragraph._element.get_or_add_pPr().append(shading_elm)

def insert_label(paragraph, label):
    label_run = paragraph.insert_paragraph_before().add_run(f'[{label}]')
    label_run.bold = True
    label_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

def annotate_images(doc):
    for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.element.xpath('.//w:drawing') or run.element.xpath('.//w:pict'):
                    insert_label(paragraph, 'This is an image')
                    set_paragraph_bg_color(paragraph, 'ADD8E6')  #Light blue for paragraphs

def process_docx(input_path, output_path):
    try:
        doc = Document(input_path)

        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                insert_label(paragraph, 'This is a Header')
                set_paragraph_bg_color(paragraph, 'FFFF00') # Yellow for headers
            elif "Equation:" in paragraph.text:
                insert_label(paragraph, 'This is an Equation')
                set_paragraph_bg_color(paragraph, '00FF00')
            elif paragraph.text.strip(): #Ensures non-empty paragraphs
                insert_label(paragraph, 'This is a Paragraph')
                set_paragraph_bg_color(paragraph, 'D3D3D3') # Light grey for paragraphs
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            set_paragraph_bg_color(paragraph, 'FFFFFF') #White for table text

        annotate_images(doc)

        doc.save(output_path)
        print(f"Annotated Document saved to {output_path}")

    except Exception as e:
        print(f"Error processing document: {e}")
    

input_path = 'C://Users//SAGAR DEEP//Desktop//Document annotate tool//MILKY_WAYed.docx'
output_path = 'annotated_document.docx'

process_docx(input_path, output_path)